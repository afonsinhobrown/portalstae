from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_documentation():
    doc = Document()
    
    # Title
    title = doc.add_heading('Documentação das Aplicações - Portal STAE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    apps = [
        {
            "nome": "1. Recursos Humanos (recursoshumanos)",
            "descricao": "É o sistema central de gestão do capital humano da instituição, responsável por centralizar toda a vida administrativa do funcionário desde o ingresso até à reforma.",
            "funcionalidades": [
                "Cadastro Digital Único: Registo completo de dados biográficos, profissionais, bancários e de contacto.",
                "Estrutura Organizacional: Mapeamento de Direções, Departamentos e Setores com definição de chefias.",
                "Gestão de Ausências: Fluxo completo de pedido, análise e aprovação de férias, licenças e faltas justificadas.",
                "Gestão Documental: Arquivo digital de documentos institucionais e histórico de auditoria por funcionário.",
                "Diretório Institucional: Listagem interna para localização rápida de colaboradores e contactos."
            ],
            "objetivos": "Modernizar e desburocratizar a gestão de pessoal, garantir a integridade dos dados e facilitar o controlo de assiduidade e planeamento de férias."
        },
        {
            "nome": "2. Gestão de Combustível (gestaocombustivel)",
            "descricao": "Solução logística dedicada ao controlo rigoroso do consumo de combustível e da manutenção da frota automóvel.",
            "funcionalidades": [
                "Registo de Viaturas: Ficha técnica detalhada de cada viatura, incluindo matrícula, marca e estado de conservação.",
                "Requisição Digital: Pedido e aprovação de combustível com geração automática de guias de abastecimento.",
                "Controlo de Fornecedores: Gestão de contratos com postos de abastecimento, monitorização de saldos e pagamentos.",
                "Prevenção e Manutenção: Agendamento de revisões mecânicas e alerta para renovação de seguros obrigatórios."
            ],
            "objetivos": "Otimizar gastos públicos com combustível, prolongar a vida útil da frota e garantir que todas as viaturas estejam em condições legais e mecânicas para as operações."
        },
        {
            "nome": "3. Gestão de Equipamentos (gestaoequipamentos)",
            "descricao": "Módulo de inventário e gestão patrimonial para o rastreio de ativos fixos e bens móveis da instituição.",
            "funcionalidades": [
                "Inventário Permanente: Classificação por categorias e tipos de equipamentos (TIC, mobiliário, etc.) com números de série e etiquetas.",
                "Fluxo de Movimentação: Registo e autorização para transferência de bens entre setores ou províncias.",
                "Gestão de Armazéns: Controlo de stock e localização física de equipamentos em depósito.",
                "Atribuição de Responsabilidade: Vinculação de cada equipamento a um setor ou funcionário específico."
            ],
            "objetivos": "Evitar a perda de património, facilitar processos de auditoria e garantir que o equipamento certo esteja disponível para quem dele necessita."
        },
        {
            "nome": "4. Credenciais (credenciais)",
            "descricao": "Sistema de gestão de identificação e acesso, essencial para garantir a segurança e a legitimidade em períodos eleitorais.",
            "funcionalidades": [
                "Gestão de Solicitantes: Registo de indivíduos, ONGs, partidos políticos e órgãos de comunicação social.",
                "Emissão Customizada: Diferentes modelos de credenciais (observadores, técnicos, imprensa) com design adaptável.",
                "Segurança com QR Code: Inclusão de códigos únicos para verificação de validade em campo (mesmo sem internet).",
                "Validação de Pedidos: Módulo de análise e aprovação de pedidos remotos com acompanhamento de status."
            ],
            "objetivos": "Controlar fidedignamente quem está autorizado a intervir ou observar atos eleitorais e prevenir fraudes de identidade."
        },
        {
            "nome": "5. DFEC (dfec)",
            "descricao": "Portal do Departamento de Formação e Educação Cívica, focado na gestão do conhecimento, formação de agentes e análise de dados.",
            "funcionalidades": [
                "Biblioteca de Manuais: Gestão de versões e publicação de manuais de procedimentos, guias operacionais e materiais de formação.",
                "Gestão de Capacitação: Planeamento de cursos, gestão de salas, formadores e inscrições de participantes.",
                "Educação Cívica: Monitorização de brigadas de sensibilização no terreno e gestão de materiais didáticos.",
                "Análise de Dados Eleitorais: Módulo analítico para estudo de resultados históricos e indicadores de abstenção."
            ],
            "objetivos": "Padronizar o conhecimento técnico em todo o país, gerir eficientemente a formação massiva de agentes eleitorais e apoiar a tomada de decisão baseada em dados estatísticos."
        }
    ]
    
    for app in apps:
        doc.add_heading(app['nome'], level=1)
        
        doc.add_heading('Descrição:', level=2)
        doc.add_paragraph(app['descricao'])
        
        doc.add_heading('Funcionalidades:', level=2)
        for func in app['funcionalidades']:
            doc.add_paragraph(func, style='List Bullet')
        
        doc.add_heading('Objetivos:', level=2)
        doc.add_paragraph(app['objetivos'])
        
        doc.add_paragraph() # Spacer
    
    doc.save('Documentacao_Portal_STAE.docx')
    print("Documento criado com sucesso: Documentacao_Portal_STAE.docx")

if __name__ == "__main__":
    create_documentation()
