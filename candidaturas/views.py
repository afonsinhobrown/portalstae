from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.db import models
from django.db.models import Count
from .models import InscricaoPartidoEleicao, ListaCandidatura, Candidato, CirculoEleitoral
from .forms import InscricaoPartidoForm, ListaCandidaturaForm, CandidatoForm, ImportarListaForm
from eleicao.models import Eleicao
from partidos.models import Partido
from rs.models import Eleitor
from django.utils import timezone
from datetime import date
from django.http import HttpResponse, JsonResponse
from docx import Document
import pandas as pd
import io


def exportar_relatorio_word(request, eleicao_id):
    """Exporta resumo demográfico para Word"""
    eleicao = get_object_or_404(Eleicao, id=eleicao_id)
    doc = Document()
    doc.add_heading(f'Relatório de Candidaturas - {eleicao.nome}', 0)
    
    candidatos = Candidato.objects.filter(lista__inscricao__eleicao=eleicao)
    doc.add_paragraph(f'Total de Candidatos: {candidatos.count()}')
    doc.add_paragraph(f'Homens: {candidatos.filter(genero="M").count()}')
    doc.add_paragraph(f'Mulheres: {candidatos.filter(genero="F").count()}')
    
    doc.add_heading('Cabeças de Lista', level=1)
    lideres = candidatos.filter(posicao=1, tipo='efetivo')
    for l in lideres:
        doc.add_paragraph(f'- {l.nome_completo} ({l.lista.inscricao.partido.sigla} | {l.lista.circulo.nome})')
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    response = HttpResponse(file_stream.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Relatorio_Candidaturas_{eleicao.ano}.docx'
    return response

def get_estatisticas_eleicao(eleicao_id):
    """Motor Central de BI: Calcula todos os KPIs para UI e PDF"""
    eleicao = get_object_or_404(Eleicao, id=eleicao_id)
    circulos = CirculoEleitoral.objects.filter(eleicao=eleicao).order_by('provincia', 'nome')
    partidos = Partido.objects.all().order_by('sigla')
    total_circulos = circulos.count()
    hoje = date.today()

    def calc_idade(dt):
        if not dt: return 0
        return hoje.year - dt.year - ((hoje.month, hoje.day) < (dt.month, dt.day))

    # 1. MAPA DE CRUZAMENTO PROVINCIAL
    mapa_por_provincia = []
    provincias_list = circulos.values_list('provincia', flat=True).distinct().order_by('provincia')
    totais_por_circulo = {c.id: 0 for c in circulos}
    
    for prov in provincias_list:
        if not prov: continue
        circs_prov = circulos.filter(provincia=prov)
        matriz_prov = []
        totais_partido_prov = {p.id: 0 for p in partidos}
        for c in circs_prov:
            presencas_circulo = []
            total_c = 0
            for p in partidos:
                existe = ListaCandidatura.objects.filter(inscricao__partido=p, circulo=c).exists()
                presencas_circulo.append({'partido_id': p.id, 'sigla': p.sigla, 'existe': existe})
                if existe:
                    total_c += 1
                    totais_partido_prov[p.id] += 1
                    totais_por_circulo[c.id] += 1
            matriz_prov.append({'circulo': c, 'status_partidos': presencas_circulo, 'total_no_circulo': total_c})
        mapa_por_provincia.append({
            'provincia': prov, 'partidos_colunas': partidos, 'matriz_linhas': matriz_prov,
            'totais_por_partido': totais_partido_prov, 'provincia_total': sum(totais_partido_prov.values())
        })
    
    grande_total = sum(totais_por_circulo.values())

    # 2. PERFORMANCE PARTIDÁRIA
    partidos_estat = []
    for p in partidos:
        listas_p = ListaCandidatura.objects.filter(inscricao__partido=p, inscricao__eleicao=eleicao)
        cands_p = Candidato.objects.filter(lista__inscricao__partido=p, lista__inscricao__eleicao=eleicao)
        
        # Cobertura Territorial
        cobertura = (listas_p.count() / total_circulos * 100) if total_circulos > 0 else 0
        
        # Conformidade Legal
        conformidade = {'apto': 0, 'irregular': 0}
        for lista in listas_p:
            if lista.verificar_conformidade()['conforme']: conformidade['apto'] += 1
            else: conformidade['irregular'] += 1

        # Dados de Gênero
        total_cands = cands_p.count()
        mulheres = cands_p.filter(genero='F').count()
        homens = total_cands - mulheres
        
        partidos_estat.append({
            'partido': p,
            'total_listas': listas_p.count(),
            'total_candidatos': total_cands,
            'cobertura': round(cobertura, 1),
            'conformidade': conformidade,
            'homens': homens,
            'mulheres': mulheres,
            'mulheres_perc': round((mulheres / total_cands * 100), 1) if total_cands > 0 else 0
        })

    # 3. DEMOGRAFIA, EXTREMOS E INTEGRIDADE
    cands_qs = Candidato.objects.filter(models.Q(lista__inscricao__eleicao=eleicao) | models.Q(inscricao_direta__eleicao=eleicao))
    
    # Função para obter extremos por QS
    def get_extremes(qs):
        validos = qs.exclude(data_nascimento=None)
        return {
            'velho_h': validos.filter(genero='M').order_by('data_nascimento').first(),
            'velho_m': validos.filter(genero='F').order_by('data_nascimento').first(),
            'novo_h': validos.filter(genero='M').order_by('-data_nascimento').first(),
            'novo_m': validos.filter(genero='F').order_by('-data_nascimento').first(),
        }

    extremes_nacionais = get_extremes(cands_qs)
    
    # Listagem Oficial de Círculos (com partidos e posições)
    listagem_oficial_circulos = []
    for c in circulos:
        listas_c = ListaCandidatura.objects.filter(circulo=c).select_related('inscricao__partido')
        # Ordenar por posicao_boletim (se existir) ou sigla
        partidos_ordenados = sorted(listas_c, key=lambda l: (l.inscricao.posicao_boletim or 999, l.inscricao.partido.sigla))
        
        listagem_oficial_circulos.append({
            'circulo': c,
            'listas': partidos_ordenados,
            'extremos': get_extremes(Candidato.objects.filter(lista__circulo=c))
        })

    # Pirâmide Geracional
    piramide = {
        'jovens': {'total': 0, 'M': 0, 'F': 0, 'faixa': '18 - 35 ANOS'},
        'adultos': {'total': 0, 'M': 0, 'F': 0, 'faixa': '36 - 60 ANOS'},
        'seniores': {'total': 0, 'M': 0, 'F': 0, 'faixa': '> 60 ANOS'}
    }
    
    for c in cands_qs.exclude(data_nascimento=None):
        idade = calc_idade(c.data_nascimento)
        gen = c.genero if c.genero in ['M', 'F'] else 'M'
        group = 'jovens' if idade <= 35 else 'adultos' if idade <= 60 else 'seniores'
        piramide[group]['total'] += 1
        piramide[group][gen] += 1

    # Auditoria de Operadores
    listas_audit = []
    for lista in ListaCandidatura.objects.filter(inscricao__eleicao=eleicao).select_related('inscricao__partido', 'circulo', 'submetido_por').order_by('-data_submissao')[:50]:
        listas_audit.append({
            'lista': lista,
            'operador': lista.submetido_por.get_full_name() or lista.submetido_por.username if lista.submetido_por else "Importação Sistema",
            'status': {'conforme': lista.validada, 'erros': [], 'alertas': []} # Versão Light para Dashboard
        })

    fraudes = cands_qs.filter(duplicado=True).count()
    
    return {
        'eleicao': eleicao,
        'mapa_por_provincia': mapa_por_provincia,
        'partidos_estat': partidos_estat,
        'listagem_oficial': listagem_oficial_circulos,
        'listas_audit': listas_audit,
        'grande_total': grande_total,
        'fraudes': fraudes,
        'demografia': {
            'total': cands_qs.count(),
            'piramide': piramide,
            'extremos': extremes_nacionais,
            'homens': cands_qs.filter(genero='M').count(),
            'mulheres': cands_qs.filter(genero='F').count(),
            'lideres_homens': cands_qs.filter(posicao=1, genero='M').count(),
            'lideres_mulheres': cands_qs.filter(posicao=1, genero='F').count(),
            'rs_confirmados': cands_qs.filter(status_eleitor='confirmado').count(),
        }
    }

def relatorios_eleicao(request, eleicao_id):
    context = get_estatisticas_eleicao(eleicao_id)
    return render(request, 'candidaturas/relatorios_painel.html', context)

def exportar_relatorio_pdf(request, eleicao_id):
    from django.template.loader import get_template
    from xhtml2pdf import pisa
    context = get_estatisticas_eleicao(eleicao_id)
    context['data_emissao'] = timezone.now()
    
    template = get_template('candidaturas/relatorio_pdf.html')
    html = template.render(context)
    result = io.BytesIO()
    pdf = pisa.pisaDocument(io.BytesIO(html.encode("UTF-8")), result)
    
    if not pdf.err:
        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=Relatorio_BI_Candidaturas.pdf'
        return response
    return HttpResponse("Erro ao gerar PDF", status=500)

def detalhe_candidato(request, candidato_id):
    """Perfil Detalhado do Candidato: Auditoria Individual"""
    candidato = get_object_or_404(Candidato, id=candidato_id)
    return render(request, 'candidaturas/detalhe_candidato.html', {'candidato': candidato})

def relatorio_global_a3(request, eleicao_id):
    """Vista de Páginas Gigantes (A3) para Monitorização Panorâmica Nacional"""
    eleicao = get_object_or_404(Eleicao, id=eleicao_id)
    circulos = CirculoEleitoral.objects.filter(eleicao=eleicao).order_by('provincia', 'nome')
    partidos = Partido.objects.all().order_by('sigla')
    
    # Gerar Matriz Nacional Completa
    matriz_nacional = []
    for c in circulos:
        presencas = []
        for p in partidos:
            existe = ListaCandidatura.objects.filter(inscricao__partido=p, circulo=c).exists()
            presencas.append({'partido': p, 'existe': existe})
        matriz_nacional.append({
            'circulo': c,
            'presencas': presencas,
            'total': sum(1 for item in presencas if item['existe'])
        })
        
    context = {
        'eleicao': eleicao,
        'partidos': partidos,
        'matriz': matriz_nacional,
    }
    return render(request, 'candidaturas/relatorio_global_a3.html', context)

def exportar_relatorio_excel(request, eleicao_id):
    """Exportação Provincial Consolidada para Auditoria"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    
    eleicao = get_object_or_404(Eleicao, id=eleicao_id)
    circulos = CirculoEleitoral.objects.filter(eleicao=eleicao).order_by('provincia', 'nome')
    partidos = Partido.objects.all().order_by('sigla')
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Consolidado Provincial"
    
    header_fill = PatternFill(start_color="1e293b", end_color="1e293b", fill_type="solid")
    white_font = Font(color="FFFFFF", bold=True)
    
    curr_row = 1
    for prov in circulos.values_list('provincia', flat=True).distinct().order_by('provincia'):
        if not prov: continue
        ws.cell(row=curr_row, column=1, value=f"PROVÍNCIA: {prov.upper()}").font = Font(bold=True, size=12)
        curr_row += 1
        
        headers = ["CÍRCULO \ PARTIDO"] + [p.sigla for p in partidos] + ["TOTAL"]
        for col, text in enumerate(headers, 1):
            cell = ws.cell(row=curr_row, column=col, value=text)
            cell.fill = header_fill
            cell.font = white_font
            
        curr_row += 1
        for c in circulos.filter(provincia=prov):
            ws.cell(row=curr_row, column=1, value=c.nome)
            total_c = 0
            for col, p in enumerate(partidos, 2):
                existe = ListaCandidatura.objects.filter(inscricao__partido=p, circulo=c).exists()
                ws.cell(row=curr_row, column=col, value="SIM" if existe else "-")
                if existe: total_c += 1
            ws.cell(row=curr_row, column=col+1, value=total_c).font = Font(bold=True)
            curr_row += 1
        curr_row += 2

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename=Auditoria_Provincial_{eleicao.ano}.xlsx'
    wb.save(response)
    return response

def gerar_dossier_lista(request, lista_id):
    """Gera Dossier Intermédio (Prova de Boletim) para Auditoria e Homologação"""
    lista = get_object_or_404(ListaCandidatura, id=lista_id)
    conformidade = lista.verificar_conformidade()
    
    # Criar um PDF ou Word Estilizado (usando Word por ser mais fácil de editar pelo STAE se necessário)
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Cabeçalho Oficial
    header = doc.add_heading('REPÚBLICA DE MOÇAMBIQUE', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SECRETARIADO TÉCNICO DE ADMINISTRAÇÃO ELEITORAL\n')
    r.bold = True
    r.font.size = Pt(12)
    p.add_run(f'DOSSIER DE HOMOLOGAÇÃO DE CANDIDATURA - {lista.inscricao.eleicao.ano}\n')
    
    doc.add_heading('1. IDENTIFICAÇÃO DA LISTA', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Partido/Coligação:'
    cells[1].text = f"{lista.inscricao.partido.sigla} - {lista.inscricao.partido.nome_completo}"
    
    cells = table.rows[1].cells
    cells[0].text = 'Círculo Eleitoral:'
    cells[1].text = f"{lista.circulo.nome} ({lista.circulo.provincia})"
    
    cells = table.rows[2].cells
    cells[0].text = 'Cargo Disputado:'
    cells[1].text = lista.cargo_disputado
    
    cells = table.rows[3].cells
    cells[0].text = 'Data de Submissão:'
    cells[1].text = timezone.now().strftime('%d/%m/%Y %H:%M')
    
    doc.add_heading('2. PARECER DE CONFORMIDADE LEGAL', level=1)
    status_text = "CONFORME (APTO PARA BOLETIM)" if conformidade['conforme'] else "IRREGULAR (REJEITADO)"
    p = doc.add_paragraph()
    r = p.add_run(f'ESTADO FINAL: {status_text}')
    r.bold = True
    
    if conformidade['erros']:
        doc.add_heading('ERROS DETECTADOS:', level=2)
        for erro in conformidade['erros']:
            doc.add_paragraph(f'- {erro}', style='List Bullet')
            
    doc.add_heading('3. ORDEM DE PRECEDÊNCIA (PROVA DE BOLETIM)', level=1)
    cand_table = doc.add_table(rows=1, cols=4)
    cand_table.style = 'Table Grid'
    hdr_cells = cand_table.rows[0].cells
    hdr_cells[0].text = 'ORD.'
    hdr_cells[1].text = 'NOME COMPLETO'
    hdr_cells[2].text = 'CARTÃO ELEITOR'
    hdr_cells[3].text = 'ESTATUTO RS'
    
    for cand in conformidade['candidatos_data']:
        row_cells = cand_table.add_row().cells
        row_cells[0].text = str(cand.posicao)
        row_cells[1].text = cand.nome_completo
        row_cells[2].text = cand.numero_eleitor or '---'
        row_cells[3].text = cand.get_status_eleitor_display()
        
    doc.add_paragraph('\n\n')
    doc.add_paragraph('__________________________________________\nAssinatura do Responsável de Mesa (STAE)')
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    response = HttpResponse(file_stream.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Dossier_Homologacao_{lista.id}.docx'
    return response

def baixar_modelo_excel(request, lista_id):
    """Gera um ficheiro Excel Dinâmico com Dropdowns de Seleção"""
    lista = get_object_or_404(ListaCandidatura, id=lista_id)
    
    # 1. Coleta de Dados para Listas Suspensas
    partidos = [f"{p.sigla} - {p.nome_completo}" for p in Partido.objects.all()]
    eleicoes = [f"{e.nome} ({e.ano})" for e in Eleicao.objects.filter(ativo=True)]
    circulos_qs = CirculoEleitoral.objects.filter(eleicao=lista.inscricao.eleicao).order_by('nome')
    
    # Colunas de Dados
    colunas = ['Nome Completo', 'BI', 'Cartão Eleitor', 'Data Nascimento (AAAA-MM-DD)', 'Género (M/F)', 'Tipo (Efetivo/Suplente)', 'Posição']
    df_main = pd.DataFrame(columns=colunas)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_main.to_excel(writer, index=False, sheet_name='Lista_Candidatos', startrow=5)
        
        workbook = writer.book
        ws = writer.sheets['Lista_Candidatos']
        
        # --- Criar Folha de Configuração (Oculta) para as Listas ---
        ws_cfg = workbook.create_sheet("Config_Sistema")
        ws_cfg.sheet_state = 'hidden'
        
        # Preencher Listas de Referência
        for i, val in enumerate(partidos, 1): ws_cfg.cell(row=i, column=1, value=val)
        for i, val in enumerate(eleicoes, 1):  ws_cfg.cell(row=i, column=2, value=val)
        for i, c in enumerate(circulos_qs, 1):
            ws_cfg.cell(row=i, column=3, value=c.nome)
            ws_cfg.cell(row=i, column=4, value=c.num_mandatos)
            
        # --- CABEÇALHO DINÂMICO ---
        from openpyxl.styles import Font, PatternFill, Alignment
        label_font = Font(bold=True, size=11)
        select_fill = PatternFill(start_color="E0F2FE", end_color="E0F2FE", fill_type="solid")
        
        ws['A1'], ws['A2'], ws['A3'], ws['A4'] = "PARTIDO:", "ELEIÇÃO:", "CÍRCULO:", "MANDATOS:"
        ws['A5'] = "SECRETARIADO TÉCNICO DE ADMINISTRAÇÃO ELEITORAL - MOÇAMBIQUE"
        
        # Valores Iniciais Pré-selecionados
        ws['B1'] = f"{lista.inscricao.partido.sigla} - {lista.inscricao.partido.nome_completo}"
        ws['B2'] = f"{lista.inscricao.eleicao.nome} ({lista.inscricao.eleicao.ano})"
        ws['B3'] = f"{lista.circulo.nome}"
        ws['B4'] = f"{lista.circulo.num_mandatos}"
        
        # Fórmula Excel para Mandatos Dinâmicos: Procura o valor de B3 na lista de círculos da folha oculta
        ws['B4'] = '=IFERROR(VLOOKUP(B3, Config_Sistema!C:D, 2, FALSE), "Selecione o Círculo")'
        
        for row in range(1, 5):
            ws[f'A{row}'].font = label_font
            ws[f'B{row}'].fill = select_fill
            ws[f'B{row}'].alignment = Alignment(horizontal='left')

        # --- VALIDAÇÕES E DROPDOWNS ---
        from openpyxl.worksheet.datavalidation import DataValidation
        
        # Dropdown de Partido (B1)
        dv_partido = DataValidation(type="list", formula1=f"Config_Sistema!$A$1:$A${len(partidos)}", allow_blank=False)
        ws.add_data_validation(dv_partido)
        dv_partido.add("B1")
        
        # Dropdown de Eleição (B2)
        dv_eleicao = DataValidation(type="list", formula1=f"Config_Sistema!$B$1:$B${len(eleicoes)}", allow_blank=False)
        ws.add_data_validation(dv_eleicao)
        dv_eleicao.add("B2")
        
        # Dropdown de Círculo (B3)
        dv_circulo = DataValidation(type="list", formula1=f"Config_Sistema!$C$1:$C${len(circulos_nomes)}", allow_blank=False)
        ws.add_data_validation(dv_circulo)
        dv_circulo.add("B3")
        
        # Dropdowns da Tabela (Género e Tipo)
        dv_gen = DataValidation(type="list", formula1='"M,F"', allow_blank=True)
        ws.add_data_validation(dv_gen)
        dv_gen.add("E7:E500")
        
        dv_tip = DataValidation(type="list", formula1='"Efetivo,Suplente"', allow_blank=True)
        ws.add_data_validation(dv_tip)
        dv_tip.add("F7:F500")
        
        # Ajustes de Layout
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 60
        for col_letter in ['C', 'D', 'E', 'F', 'G']:
            ws.column_dimensions[col_letter].width = 20

    output.seek(0)
    response = HttpResponse(output.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename=Submissao_Candidatos_STAE.xlsx'
    return response

def baixar_modelo_vazio(request, inscricao_id):
    """Gera um modelo Excel para um partido que ainda não tem listas criadas"""
    insc = get_object_or_404(InscricaoPartidoEleicao, id=inscricao_id)
    
    # 1. Coleta de Dados
    partidos = [f"{insc.partido.sigla} - {insc.partido.nome_completo}"]
    eleicoes = [f"{insc.eleicao.nome} ({insc.eleicao.ano})"]
    
    # Se for autárquica, os círculos são os do sistema para aquela eleição
    circulos_qs = CirculoEleitoral.objects.filter(eleicao=insc.eleicao, ativo=True).order_by('nome')
    # Se for provincial, filtramos apenas os círculos daquela província
    if insc.scope_provincia:
        circulos_qs = circulos_qs.filter(provincia=insc.scope_provincia.nome)
        
    circulos_nomes = [c.nome for c in circulos_qs]
    
    colunas = ['Nome Completo', 'BI', 'Cartão Eleitor', 'Data Nascimento (AAAA-MM-DD)', 'Género (M/F)', 'Tipo (Efetivo/Suplente)', 'Posição']
    df_main = pd.DataFrame(columns=colunas)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_main.to_excel(writer, index=False, sheet_name='Lista_Candidatos', startrow=5)
        workbook = writer.book
        ws = writer.sheets['Lista_Candidatos']
        
        ws_cfg = workbook.create_sheet("Config_Sistema")
        ws_cfg.sheet_state = 'hidden'
        for i, val in enumerate(partidos, 1): ws_cfg.cell(row=i, column=1, value=val)
        for i, val in enumerate(eleicoes, 1):  ws_cfg.cell(row=i, column=2, value=val)
        for i, c in enumerate(circulos_qs, 1):
            ws_cfg.cell(row=i, column=3, value=c.nome)
            ws_cfg.cell(row=i, column=4, value=c.num_mandatos)
            
        from openpyxl.styles import Font, PatternFill, Alignment
        label_font = Font(bold=True, size=11)
        select_fill = PatternFill(start_color="E0F2FE", end_color="E0F2FE", fill_type="solid")
        
        ws['A1'], ws['A2'], ws['A3'], ws['A4'] = "PARTIDO:", "ELEIÇÃO:", "CÍRCULO:", "MANDATOS:"
        ws['A5'] = "SECRETARIADO TÉCNICO DE ADMINISTRAÇÃO ELEITORAL - MOÇAMBIQUE (Template Geral)"
        
        ws['B1'] = partidos[0]
        ws['B2'] = eleicoes[0]
        ws['B3'] = "Selecione o Círculo"
        ws['B4'] = '=IFERROR(VLOOKUP(B3, Config_Sistema!C:D, 2, FALSE), "---")'
        
        for row in range(1, 5):
            ws[f'A{row}'].font = label_font
            ws[f'B{row}'].fill = select_fill
            
        from openpyxl.worksheet.datavalidation import DataValidation
        ws.add_data_validation(DataValidation(type="list", formula1=f"Config_Sistema!$C$1:$C${len(circulos_nomes)}", allow_blank=False)).add("B3")
        
        dv_gen = DataValidation(type="list", formula1='"M,F"', allow_blank=True)
        ws.add_data_validation(dv_gen); dv_gen.add("E7:E500")
        dv_tip = DataValidation(type="list", formula1='"Efetivo,Suplente"', allow_blank=True)
        ws.add_data_validation(dv_tip); dv_tip.add("F7:F500")

        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 60

    output.seek(0)
    response = HttpResponse(output.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename=Template_Submissao_{insc.partido.sigla}.xlsx'
    return response

def baixar_dados_teste_excel(request):
    """Gera um ficheiro Excel de Simulação (Maxixe - 40 Mandatos) com Dropdowns e cabeçalho dinâmico"""
    # 1. Dados de Simulação (40 Efetivos + 3 Suplentes)
    nomes = [f"Candidato Simulação {i}" for i in range(1, 44)]
    bis = [f"{10000000+i}A" for i in range(1, 44)]
    eleitores = [f"{20202020+i}" for i in range(1, 44)]
    nascimentos = ["1992-05-10"] * 43
    generos = ["M" if i%2==0 else "F" for i in range(43)]
    tipos = ["Efetivo"] * 40 + ["Suplente"] * 3
    posicoes = list(range(1, 41)) + list(range(1, 4))
    
    df_dados = pd.DataFrame({
        'Nome Completo': nomes,
        'BI': bis,
        'Cartão Eleitor': eleitores,
        'Data Nascimento (AAAA-MM-DD)': nascimentos,
        'Género (M/F)': generos,
        'Tipo (Efetivo/Suplente)': tipos,
        'Posição': posicoes
    })

    # 2. Dados de Sistema para os Dropdowns
    partidos = [f"{p.sigla} - {p.nome_completo}" for p in Partido.objects.all()]
    eleicoes = [f"{e.nome} ({e.ano})" for e in Eleicao.objects.filter(ativo=True)]
    circulos_qs = CirculoEleitoral.objects.all().order_by('nome')
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_dados.to_excel(writer, index=False, sheet_name='Simulacao_STAE', startrow=5)
        
        workbook = writer.book
        ws = writer.sheets['Simulacao_STAE']
        
        # --- Folha de Configuração ---
        ws_cfg = workbook.create_sheet("Config_Sistema")
        ws_cfg.sheet_state = 'hidden'
        for i, val in enumerate(partidos, 1): ws_cfg.cell(row=i, column=1, value=val)
        for i, val in enumerate(eleicoes, 1):  ws_cfg.cell(row=i, column=2, value=val)
        for i, c in enumerate(circulos_qs, 1):
            ws_cfg.cell(row=i, column=3, value=c.nome)
            ws_cfg.cell(row=i, column=4, value=c.num_mandatos)
            
        # --- CABEÇALHO DINÂMICO ---
        from openpyxl.styles import Font, PatternFill, Alignment
        label_font = Font(bold=True, size=11)
        select_fill = PatternFill(start_color="E0F2FE", end_color="E0F2FE", fill_type="solid")
        
        ws['A1'], ws['A2'], ws['A3'], ws['A4'] = "PARTIDO:", "ELEIÇÃO:", "CÍRCULO:", "MANDATOS:"
        ws['A5'] = "SECRETARIADO TÉCNICO DE ADMINISTRAÇÃO ELEITORAL - MOÇAMBIQUE"
        
        # Pre-seleção para Simulação
        ws['B1'] = partidos[0] if partidos else "Selecione o Partido"
        ws['B2'] = eleicoes[0] if eleicoes else "Selecione a Eleição"
        ws['B3'] = "CIDADE DE MAXIXE"
        ws['B4'] = '=IFERROR(VLOOKUP(B3, Config_Sistema!C:D, 2, FALSE), "Selecione")'
        
        for row in range(1, 5):
            ws[f'A{row}'].font = label_font
            ws[f'B{row}'].fill = select_fill
            ws[f'B{row}'].alignment = Alignment(horizontal='left')

        # --- DROP-DOWNS ---
        from openpyxl.worksheet.datavalidation import DataValidation
        ws.add_data_validation(DataValidation(type="list", formula1=f"Config_Sistema!$A$1:$A${len(partidos)}", allow_blank=False)).add("B1")
        ws.add_data_validation(DataValidation(type="list", formula1=f"Config_Sistema!$B$1:$B${len(eleicoes)}", allow_blank=False)).add("B2")
        ws.add_data_validation(DataValidation(type="list", formula1=f"Config_Sistema!$C$1:$C${len(circulos_qs)}", allow_blank=False)).add("B3")
        
        # Dropdowns da Tabela
        dv_gen = DataValidation(type="list", formula1='"M,F"', allow_blank=True)
        ws.add_data_validation(dv_gen); dv_gen.add("E7:E500")
        dv_tip = DataValidation(type="list", formula1='"Efetivo,Suplente"', allow_blank=True)
        ws.add_data_validation(dv_tip); dv_tip.add("F7:F500")

        # Ajuste de Layout
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 60
        for col in ['C', 'D', 'E', 'F', 'G']: ws.column_dimensions[col].width = 20

    output.seek(0)
    response = HttpResponse(output.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=SIMULACAO_40_MANDATOS_MAXIXE.xlsx'
    return response

def dashboard(request):
    eleicoes = Eleicao.objects.filter(ativo=True).order_by('-ano')
    inscricoes = InscricaoPartidoEleicao.objects.all().select_related('partido', 'eleicao', 'scope_provincia', 'scope_circulo')
    candidatos = Candidato.objects.all()
    
    # --- AUDITORIA DE INTEGRIDADE (Anti-Fraude) ---
    # Otimizado: Só atualizamos se houver incidência de novos dados ou via rotina manual
    duplicados_count = Candidato.objects.filter(duplicado=True).count()

    # --- ESTATÍSTICAS DEMOGRÁFICAS ---
    stats_gen = candidatos.aggregate(
        homens=Count('id', filter=models.Q(genero='M')),
        mulheres=Count('id', filter=models.Q(genero='F'))
    )
    
    hoje = date.today()
    jovens, adultos, seniores = 0, 0, 0
    # Otimização: Só processamos quem tem data
    candidatos_com_data = candidatos.exclude(data_nascimento__isnull=True).only('data_nascimento')
    for c in candidatos_com_data:
        idade = hoje.year - c.data_nascimento.year - ((hoje.month, hoje.day) < (c.data_nascimento.month, c.data_nascimento.day))
        if idade <= 35: jovens += 1
        elif idade <= 60: adultos += 1
        else: seniores += 1

    # Alcance de Partidos (Geografia Dinâmica)
    # Procuramos por inscrições na eleição ativa
    eleicao_ativa = eleicoes.first()
    total_circulos_sistema = CirculoEleitoral.objects.filter(eleicao=eleicao_ativa).count() if eleicao_ativa else 1
    if total_circulos_sistema == 0: total_circulos_sistema = 1
    
    alcance_partidos = InscricaoPartidoEleicao.objects.filter(eleicao=eleicao_ativa).values(
        'partido__sigla', 'partido__nome_completo'
    ).annotate(
        num_circulos=Count('listas', filter=models.Q(listas__candidatos__isnull=False), distinct=True)
    ).order_by('-num_circulos')

    # Mapeamento de âmbitos
    ambitos_map = {}
    for ins in inscricoes:
        sigla = ins.partido.sigla
        if sigla not in ambitos_map: ambitos_map[sigla] = []
        texto_ambito = "Global"
        if ins.scope_provincia: texto_ambito = ins.scope_provincia.nome
        elif ins.scope_circulo: texto_ambito = ins.scope_circulo.nome
        ambitos_map[sigla].append(texto_ambito)

    for a in alcance_partidos:
        a['ambitos'] = ambitos_map.get(a['partido__sigla'], ["Nacional"])

    # Detalhar listas para conformidade
    listas_info = ListaCandidatura.objects.filter(inscricao__eleicao=eleicao_ativa).select_related('circulo')
    conformidade_map = {}
    for l in listas_info:
        conformidade_map[l.id] = l.verificar_conformidade()

    context = {
        'eleicoes': eleicoes,
        'inscricoes': inscricoes,
        'total_partidos': Partido.objects.count(),
        'total_candidatos': candidatos.count(),
        'total_circulos_sistema': total_circulos_sistema,
        'duplicados_detectados': duplicados_count,
        'conformidade_map': conformidade_map,
        'stats': {
            'homens': stats_gen['homens'], 'mulheres': stats_gen['mulheres'],
            'jovens': jovens, 'adultos': adultos, 'seniores': seniores,
            'alcance': alcance_partidos
        }
    }
    return render(request, 'candidaturas/dashboard.html', context)

def baixar_ficheiro_teste(request, tipo):
    """Serve os ficheiros de teste gerados para o utilizador"""
    files = {
        'maputo_p1': 'Candidaturas_P1_Maputo.xlsx',
        'maputo_p2': 'Candidaturas_P2_Maputo.xlsx',
        'xaixai_p2': 'Candidaturas_P2_XaiXai.xlsx',
        'maxixe_p1': 'Teste_P1_Maxixe.xlsx' # Mantendo este se existir
    }
    filename = files.get(tipo)
    if not filename or not os.path.exists(filename):
        messages.error(request, "Ficheiro de teste não encontrado.")
        return redirect('candidaturas:dashboard')
    
    with open(filename, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename={filename}'
        return response

def detalhe_candidato(request, candidato_id):
    """Ficha detalhada do candidato para auditoria e verificação de BI"""
    candidato = get_object_or_404(Candidato, id=candidato_id)
    
    # Calcular idade para o perfil
    idade = "N/A"
    if candidato.data_nascimento:
        hoje = date.today()
        idade = hoje.year - candidato.data_nascimento.year - ((hoje.month, hoje.day) < (candidato.data_nascimento.month, candidato.data_nascimento.day))
    
    context = {
        'candidato': candidato,
        'idade': idade,
        'lista': candidato.lista
    }
    return render(request, 'candidaturas/detalhe_candidato.html', context)

def registrar_presidente(request):
    """Funcionalidade específica para candidatos presidenciais"""
    if request.method == 'POST':
        form = CandidatoForm(request.POST, request.FILES)
        if form.is_valid():
            cand = form.save(commit=False)
            cand.categoria = 'presidencial'
            cand.posicao = 1
            cand.tipo = 'efetivo'
            cand.save()
            messages.success(request, f"Candidato Presidencial {cand.nome_completo} registado com sucesso!")
            return redirect('candidaturas:dashboard')
    else:
        form = CandidatoForm(initial={'categoria': 'presidencial', 'posicao': 1})
    
    return render(request, 'candidaturas/form_candidato.html', {
        'form': form, 
        'titulo': 'Registo Oficial de Candidato Presidencial',
        'subtitulo': 'Esta candidatura concorre a nível nacional'
    })

def processar_submissao_excel(request, inscricao_id):
    """Central de Importação Unificada: Resiliência Máxima e Feedback Constante"""
    insc = get_object_or_404(InscricaoPartidoEleicao, id=inscricao_id)
    
    if request.method == 'POST':
        file = request.FILES.get('arquivo_excel')
        if not file:
            messages.error(request, 'Nenhum ficheiro selecionado.')
            return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao_id)

        try:
            # 1. Leitura Inicial e Localização do Círculo (Busca em Matriz)
            file.seek(0)
            df_header = pd.read_excel(file, nrows=10, header=None) # Lemos as primeiras 10 linhas para busca
            
            circulo_nome_raw = None
            partido_excel = None
            
            # Procurar por palavras-chave na coluna A e valor na coluna B
            for i, row in df_header.iterrows():
                label = str(row[0]).upper() if not pd.isna(row[0]) else ""
                valor = str(row[1]).strip() if not pd.isna(row[1]) else ""
                
                if "CIRCULO" in label or "CÍRCULO" in label:
                    circulo_nome_raw = valor
                if "PARTIDO" in label:
                    partido_excel = valor

            if not circulo_nome_raw:
                # Fallback: Tentar B3 especificamente como última tentativa
                circulo_nome_raw = str(df_header.iloc[2, 1]).strip() if df_header.shape[0] > 2 else ""

            if not circulo_nome_raw or circulo_nome_raw == "nan":
                messages.error(request, "ERRO: Não foi possível localizar o nome do Círculo Eleitoral no cabeçalho do Excel. Verifique se o nome está na Célula B3.")
                return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao_id)

            # Normalização robusta
            import re
            def norm(t): return re.sub(r'\s+', ' ', str(t)).strip().upper()
            
            circulo_nome_norm = norm(circulo_nome_raw)
            circulo = CirculoEleitoral.objects.filter(nome__iexact=circulo_nome_raw, eleicao=insc.eleicao).first()
            if not circulo:
                todos_ativos = CirculoEleitoral.objects.filter(eleicao=insc.eleicao)
                # 1. Tentativa por normalização total
                for c in todos_ativos:
                    if norm(c.nome) == circulo_nome_norm:
                        circulo = c
                        break
                
                # 2. Tentativa por "Contém" (Fuzzy Match para casos como "Maxixe" vs "CIDADE DE MAXIXE")
                if not circulo:
                    for c in todos_ativos:
                        db_norm = norm(c.nome)
                        if circulo_nome_norm in db_norm or db_norm in circulo_nome_norm:
                            circulo = c
                            break

            if not circulo:
                messages.error(request, f"BLOQUEIO TERRITORIAL: O círculo '{circulo_nome_raw}' não foi reconhecido pelo STAE para esta eleição.")
                return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao_id)

            # 2. Preparação da Lista
            lista, created = ListaCandidatura.objects.get_or_create(
                inscricao=insc, circulo=circulo,
                defaults={
                    'cargo_disputado': 'Assembleia Municipal' if insc.eleicao.tipo == 'autarquica' else 'Assembleia da República',
                    'submetido_por': request.user if request.user.is_authenticated else None
                }
            )
            if not created:
                lista.submetido_por = request.user if request.user.is_authenticated else None
                lista.save()

            # 3. Processamento Massivo de Candidatos
            file.seek(0)
            df = pd.read_excel(file, skiprows=5)
            df.columns = [str(c).strip().lower() for c in df.columns]

            # Mapeamento Ultra-Flexível
            cols = {
                'n': next((c for c in df.columns if any(k in c for k in ['nome', 'candidato', 'completo'])), None),
                'b': next((c for c in df.columns if 'bi' in c or 'ident' in c), None),
                'e': next((c for c in df.columns if any(k in c for k in ['eleitor', 'cart', 'nuit'])), None),
                'd': next((c for c in df.columns if any(k in c for k in ['nasc', 'data', 'birth'])), None),
                't': next((c for c in df.columns if 'tipo' in c or 'efet' in c), None),
                'p': next((c for c in df.columns if 'posi' in c or 'ordem' in c), None),
                'g': next((c for c in df.columns if 'g' in c and ('ner' in c or 'sexo' in c)), None)
            }

            if not cols['n']:
                messages.error(request, "ERRO FORMATO: Não localizei a coluna 'Nome Completo'. O preenchimento deve começar na linha 7.")
                return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao_id)

            lista.candidatos.all().delete()
            criados = 0
            erros_linha = 0

            def clean_val(v):
                if pd.isna(v): return ""
                try: return str(int(float(v)))
                except: return str(v).strip()

            candidatos_para_criar = []
            contagem_efetivos = 0
            
            for idx, row in df.iterrows():
                v_nome = row.get(cols['n'])
                if pd.isna(v_nome) or str(v_nome).strip() == "": continue
                
                try:
                    tipo_c = str(row.get(cols['t'], 'Efetivo')).strip().lower()
                    if 'supl' in tipo_c: tipo_c = 'suplente'
                    else: 
                        tipo_c = 'efetivo'
                        contagem_efetivos += 1
                    
                    # Tratar Data
                    v_nasc = row.get(cols['d'])
                    data_nasc = None
                    if not pd.isna(v_nasc):
                        try: data_nasc = pd.to_datetime(v_nasc).date()
                        except: pass

                    # Tratar Género
                    v_gen = str(row.get(cols['g'], 'M')).strip().upper()
                    genero = 'M'
                    if v_gen.startswith('F') or 'MULHER' in v_gen or 'FEM' in v_gen:
                        genero = 'F'

                    candidatos_para_criar.append(Candidato(
                        lista=lista,
                        nome_completo=str(v_nome).strip().upper(),
                        bi_numero=clean_val(row.get(cols['b'])),
                        numero_eleitor=clean_val(row.get(cols['e'])),
                        data_nascimento=data_nasc,
                        posicao=int(float(row.get(cols['p'], (criados + 1)))),
                        tipo=tipo_c,
                        genero=genero,
                        categoria=lista.cargo_disputado
                    ))
                    criados += 1
                except Exception as ex:
                    erros_linha += 1
                    print(f"Erro na linha {idx+7}: {str(ex)}")

            # --- VALIDAÇÃO DE CONFORMIDADE LEGAL (MANDATOS) ---
            mandatos_esperados = circulo.num_mandatos
            if contagem_efetivos != mandatos_esperados:
                messages.error(request, f"REJEITADO: O círculo '{circulo.nome}' exige exatamente {mandatos_esperados} candidatos efetivos. O seu ficheiro contém {contagem_efetivos}. Corrija a lista e tente novamente.")
                return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao_id)

            # Só guardamos se passar a validação de mandatos
            lista.candidatos.all().delete()
            Candidato.objects.bulk_create(candidatos_para_criar)

            if criados > 0:
                messages.success(request, f"SUBMISSÃO BEM-SUCEDIDA: {criados} candidatos registados em {circulo.nome}.")
                if erros_linha > 0:
                    messages.warning(request, f"Aviso: {erros_linha} linhas foram ignoradas por erros de formatação.")
                return render(request, 'candidaturas/resultado_importacao.html', {
                    'inscricao': insc, 'circulo': circulo, 'total': criados, 'lista': lista
                })
            else:
                messages.error(request, "O ficheiro foi lido, mas não foram encontrados candidatos válidos nas colunas.")
                return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao_id)

        except Exception as e:
            import traceback
            messages.error(request, f"ERRO CRÍTICO DE PROCESSAMENTO: {str(e)}")
            print(traceback.format_exc())
            return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao_id)

    return render(request, 'candidaturas/central_importacao.html', {'inscricao': insc})

def registrar_partido_eleicao(request):
    """Registo respeitando as regras geográficas (Escopo Autárquico vs Provincial)"""
    if request.method == 'POST':
        form = InscricaoPartidoForm(request.POST)
        if form.is_valid():
            instancia = form.save()
            messages.success(request, f"Inscrição de {instancia.partido.sigla} confirmada com sucesso para as eleições {instancia.eleicao.ano}.")
            return redirect('candidaturas:dashboard')
    else:
        form = InscricaoPartidoForm()
        
    # Filtragem Proativa de Círculos para evitar poluição visual de outras províncias 
    # se a província já estiver pré-selecionada no POST por erro de validação
    provincia_id = request.POST.get('scope_provincia') if request.method == 'POST' else None
    if provincia_id:
        from dfec.models.election_analytics import Provincia
        try:
            prov = Provincia.objects.get(id=provincia_id)
            form.fields['scope_circulo'].queryset = CirculoEleitoral.objects.filter(provincia__iexact=prov.nome.strip(), ativo=True).order_by('nome')
        except Provincia.DoesNotExist:
            pass
            
    return render(request, 'candidaturas/form_inscricao.html', {'form': form})

def gerenciar_listas(request, inscricao_id):
    inscricao = get_object_or_404(InscricaoPartidoEleicao, id=inscricao_id)
    listas = inscricao.listas.all().annotate(num_candidatos=Count('candidatos'))
    
    if request.method == 'POST':
        form = ListaCandidaturaForm(request.POST)
        if form.is_valid():
            lista = form.save(commit=False)
            lista.inscricao = inscricao
            lista.save()
            messages.success(request, f"Nova lista criada para o círculo {lista.circulo.nome}")
            return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao_id)
    else:
        form = ListaCandidaturaForm(initial={'inscricao': inscricao})
        # Filtragem Territorial Inteligente para Inscrições Provinciais/Legislativas
        if inscricao.scope_provincia:
            form.fields['circulo'].queryset = CirculoEleitoral.objects.filter(
                provincia__iexact=inscricao.scope_provincia.nome.strip(),
                eleicao=inscricao.eleicao,
                ativo=True
            ).order_by('nome')
        elif inscricao.scope_circulo:
            # Se já está restrito a um círculo (Autárquicas), limitamos o dropdown
            form.fields['circulo'].queryset = CirculoEleitoral.objects.filter(id=inscricao.scope_circulo.id)
            form.initial['circulo'] = inscricao.scope_circulo
            
    return render(request, 'candidaturas/listas_partido.html', {
        'inscricao': inscricao, 'listas': listas, 'form': form
    })

def detalhe_lista(request, lista_id):
    lista = get_object_or_404(ListaCandidatura, id=lista_id)
    candidatos = lista.candidatos.all().order_by('tipo', 'posicao')
    if request.method == 'POST':
        form = CandidatoForm(request.POST)
        if form.is_valid():
            candidato = form.save(commit=False)
            candidato.lista = lista
            eleitor = Eleitor.objects.filter(numero_cartao=candidato.numero_eleitor).first()
            candidato.status_eleitor = 'confirmado' if eleitor else 'pendente'
            candidato.save()
            messages.success(request, "Candidato adicionado à lista!")
            return redirect('candidaturas:detalhe_lista', lista_id=lista_id)
    else:
        form = CandidatoForm()
    return render(request, 'candidaturas/detalhe_lista.html', {
        'lista': lista, 'candidatos': candidatos, 'form': form
    })

def remover_lista(request, lista_id):
    lista = get_object_or_404(ListaCandidatura, id=lista_id)
    inscricao, nome_circulo = lista.inscricao, lista.circulo.nome
    is_autarquica = inscricao.eleicao.tipo == 'autarquica'
    lista.delete()
    if is_autarquica:
        inscricao.delete()
        messages.warning(request, f"Registo ANULADO: Inscrição no círculo {nome_circulo} removida.")
        return redirect('candidaturas:dashboard')
    messages.success(request, f"Lista {nome_circulo} removida.")
    return redirect('candidaturas:gerenciar_listas', inscricao_id=inscricao.id)

def remover_candidato(request, candidato_id):
    candidato = get_object_or_404(Candidato, id=candidato_id)
    lista_id, nome = (candidato.lista.id if candidato.lista else None), candidato.nome_completo
    candidato.delete()
    messages.success(request, f"Candidato {nome} removido.")
    return redirect('candidaturas:detalhe_lista', lista_id=lista_id) if lista_id else redirect('candidaturas:dashboard')

def reorganizar_posicoes(request, lista_id):
    if request.method == 'POST':
        posicoes, ids = request.POST.getlist('posicoes'), request.POST.getlist('candidato_ids')
        for cid, pos in zip(ids, posicoes):
            Candidato.objects.filter(id=cid).update(posicao=int(pos))
        messages.success(request, "Posições reorganizadas!")
    return redirect('candidaturas:detalhe_lista', lista_id=lista_id)

def remover_inscricao(request, inscricao_id):
    """Remove integralmente uma inscrição de partido e as suas listas/candidatos"""
    inscricao = get_object_or_404(InscricaoPartidoEleicao, id=inscricao_id)
    sigla = inscricao.partido.sigla
    inscricao.delete()
    messages.warning(request, f"Registo de Inscrição ({sigla}) e listas associadas foram removidos com sucesso.")
    return redirect('candidaturas:dashboard')

def api_get_circulos(request):
    """API para Dropdown Dependente: Filtra círculos por Província e Eleição"""
    provincia_id = request.GET.get('provincia_id')
    eleicao_id = request.GET.get('eleicao_id')
    
    if not provincia_id:
        return JsonResponse([], safe=False)
        
    from dfec.models.election_analytics import Provincia
    provincia = get_object_or_404(Provincia, id=provincia_id)
    
    # Ajuste Case-Insensitive para Moçambique (Normalmente Províncias estão em MAIÚSCULAS no DB)
    circulos = CirculoEleitoral.objects.filter(provincia__iexact=provincia.nome.strip(), ativo=True)
    
    if eleicao_id:
        circulos = circulos.filter(eleicao_id=eleicao_id)
        
    data = [{'id': c.id, 'nome': f"{c.nome} (Cód: {c.codigo})"} for c in circulos.order_by('nome')]
    return JsonResponse(data, safe=False)
